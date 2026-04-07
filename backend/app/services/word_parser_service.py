import re
from dataclasses import dataclass, field

from docx import Document as DocxDocument
from app.core.logging import get_logger

logger = get_logger(__name__)


@dataclass
class HeadingInfo:
    paragraph_index: int
    level: int
    text: str


@dataclass
class ParseResult:
    markdown: str
    word_count: int
    heading_count: int
    paragraph_count: int
    heading_levels: list[HeadingInfo] = field(default_factory=list)


def parse_docx(file_path: str) -> ParseResult:
    """Parse a .docx file into Markdown with real heading levels and structured tables.

    Returns a ParseResult with the full Markdown text and metadata.
    Headings come from Word styles (Heading 1/2/3/4), not regex detection.
    Tables are converted to Markdown table syntax with headers.
    """
    doc = DocxDocument(file_path)

    # Build O(1) lookup maps (avoids O(n*m) on large docs)
    para_map = {para._element: para for para in doc.paragraphs}
    # Include nested tables: doc.tables only has top-level, iterate all tbl elements
    table_map = {}
    for table in doc.tables:
        table_map[table._element] = table
    # Also find nested tables via xpath (tables inside table cells)
    for nested_tbl in doc.element.body.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tbl"):
        if nested_tbl not in table_map:
            from docx.table import Table
            table_map[nested_tbl] = Table(nested_tbl, doc)

    markdown_parts: list[str] = []
    headings: list[HeadingInfo] = []
    para_index = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "tbl":
            table = table_map.get(element)
            if table is not None:
                md_table = _table_to_markdown(table)
                if md_table:
                    markdown_parts.append(md_table)
            para_index += 1

        elif tag == "p":
            para = para_map.get(element)
            if para is not None:
                text = para.text.strip()
                if not text:
                    para_index += 1
                    continue

                style_name = para.style.name if para.style else ""

                # Convert Word heading styles to Markdown
                if style_name.startswith("Heading"):
                    level = _extract_heading_level(style_name)
                    prefix = "#" * level
                    markdown_parts.append(f"{prefix} {text}")
                    headings.append(HeadingInfo(
                        paragraph_index=para_index,
                        level=level,
                        text=text,
                    ))
                elif _is_all_bold(para):
                    # Bold paragraphs that aren't headings — treat as potential sub-headers
                    markdown_parts.append(f"**{text}**")
                else:
                    markdown_parts.append(text)

                para_index += 1

    full_markdown = "\n\n".join(markdown_parts)
    word_count = len(full_markdown.split())

    logger.info(
        "Parsed .docx: %d paragraphs, %d headings, %d words",
        para_index, len(headings), word_count,
    )

    return ParseResult(
        markdown=full_markdown,
        word_count=word_count,
        heading_count=len(headings),
        paragraph_count=para_index,
        heading_levels=headings,
    )


def _extract_heading_level(style_name: str) -> int:
    """Extract heading level from style name like 'Heading 1', 'Heading 2', etc.

    Handles: 'Heading 1', 'Heading 2', 'heading 3', 'Titlu 1' (Romanian Word).
    Ignores trailing numbers in custom styles like 'Heading 1 - Copy 2'.
    """
    # Match "Heading N" or "Titlu N" at the start
    match = re.match(r"(?:heading|titlu)\s+(\d)", style_name, re.IGNORECASE)
    if match:
        return min(int(match.group(1)), 6)
    return 1


def _is_all_bold(para) -> bool:
    """Check if all non-empty runs in a paragraph are bold."""
    runs = [r for r in para.runs if r.text.strip()]
    return bool(runs) and all(r.bold for r in runs)



def _table_to_markdown(table) -> str:
    """Convert a python-docx Table to Markdown table syntax."""
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            text = cell.text.strip().replace("\n", " ").replace("|", "\\|")
            cells.append(text)
        rows.append("| " + " | ".join(cells) + " |")

    if len(rows) < 1:
        return ""

    # Add header separator after first row
    if len(rows) >= 2:
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.insert(1, separator)
    else:
        # Single row table — add separator anyway
        col_count = len(table.rows[0].cells)
        separator = "| " + " | ".join(["---"] * col_count) + " |"
        rows.append(separator)

    return "\n".join(rows)
